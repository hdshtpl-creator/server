"""
ingest.py — Đưa tài liệu vào kho: đọc → OCR → chia đoạn → vector → lưu.
Chạy: python -m app.ingest data/raw [doc_type]
"""
import csv
import hashlib
import io
import json
import os
import re
import sys
import zipfile
from dataclasses import dataclass, field
from datetime import date, datetime, time
from pathlib import Path
from xml.etree import ElementTree

from app import db
from app.models import embed, summarize

# Kích thước ĐOẠN MỤC TIÊU, không phải kích thước cố định. `chunk_generic` cắt
# theo tiêu đề mục và câu trọn vẹn, nên đoạn thực tế ngắn/dài quanh con số này
# tuỳ nội dung — mục ngắn giữ nguyên một đoạn, mục dài tách ở chỗ chuyển ý.
CHUNK_WORDS = 320
# Số CÂU chồng lấn giữa hai đoạn liền nhau (không phải số từ): câu đầu đoạn sau
# nhắc lại ý cuối đoạn trước để đọc rời từng đoạn vẫn không đứt mạch.
CHUNK_OVERLAP_UNITS = 1
SUPPORTED_EXTENSIONS = frozenset({".txt", ".md", ".docx", ".doc", ".pdf", ".xlsx", ".csv"})


def _positive_int_env(name, default, minimum=1):
    try:
        value = int(os.getenv(name, str(default)))
    except (TypeError, ValueError):
        return default
    return value if value >= minimum else default


# Các giới hạn này bảo vệ worker khỏi file lỗi/zip bomb và bảng tính quá lớn.
# Có thể tăng qua .env sau khi đã đo RAM trên máy chủ.
MAX_SOURCE_BYTES = _positive_int_env("INGEST_MAX_SOURCE_BYTES", 100 * 1024 * 1024)
MAX_SPREADSHEET_ROWS = _positive_int_env("INGEST_MAX_SPREADSHEET_ROWS", 20_000)
MAX_SPREADSHEET_COLUMNS = _positive_int_env("INGEST_MAX_SPREADSHEET_COLUMNS", 256)
MAX_EXTRACTED_CHARS = _positive_int_env("INGEST_MAX_EXTRACTED_CHARS", 3_000_000)
MAX_CELL_CHARS = _positive_int_env("INGEST_MAX_CELL_CHARS", 4_000)
MAX_OCR_PAGES = _positive_int_env("INGEST_MAX_OCR_PAGES", 200)
MAX_ARCHIVE_UNCOMPRESSED_BYTES = _positive_int_env(
    "INGEST_MAX_ARCHIVE_UNCOMPRESSED_BYTES", 250 * 1024 * 1024)


class ExtractionError(RuntimeError):
    """Lỗi trích xuất có mã ổn định để log/dashboard hiển thị và test được."""

    def __init__(self, code: str, message: str, hint: str = ""):
        super().__init__(message)
        self.code = code
        self.message = message
        self.hint = hint

    def __str__(self):
        return f"{self.message}{' — ' + self.hint if self.hint else ''}"

    def as_dict(self):
        return {"code": self.code, "message": self.message, "hint": self.hint}


@dataclass
class ExtractionResult:
    """Kết quả trích xuất giàu metadata; ``extract_text`` cũ vẫn trả về chuỗi."""

    text: str
    format: str
    method: str
    warnings: list[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    @property
    def status(self):
        return "warning" if self.warnings else "ok"


@dataclass
class ChunkPiece:
    """Một đoạn dùng để embedding, kèm vị trí có thể mở/hiển thị khi trích dẫn."""

    content: str
    page_number: int | None = None
    section_title: str | None = None
    source_locator: str = "document"


@dataclass
class _SourceSegment:
    content: str
    page_number: int | None = None
    section_title: str | None = None
    source_locator: str = "document"


def safe_path_component(value: str, max_length=180) -> str:
    """Biến tên Drive thành một thành phần đường dẫn, không cho phép thoát thư mục đích."""
    original = str(value or "")
    cleaned = re.sub(r'[\x00-\x1f<>:"/\\|?*]', "_", original).strip(" .")
    if not cleaned or cleaned in {".", ".."}:
        cleaned = "tep"
    reserved = {"CON", "PRN", "AUX", "NUL", *(f"COM{i}" for i in range(1, 10)),
                *(f"LPT{i}" for i in range(1, 10))}
    if Path(cleaned).stem.upper() in reserved:
        cleaned = "_" + cleaned
    if cleaned != original:
        suffix = Path(cleaned).suffix
        stem = cleaned[:-len(suffix)] if suffix else cleaned
        digest = hashlib.sha256(original.encode("utf-8", errors="replace")).hexdigest()[:8]
        cleaned = f"{stem}_{digest}{suffix}"
    if len(cleaned) > max_length:
        suffix = Path(cleaned).suffix[:20]
        digest = hashlib.sha256(original.encode("utf-8", errors="replace")).hexdigest()[:8]
        cleaned = f"{cleaned[:max_length - len(suffix) - 10]}_{digest}{suffix}"
    return cleaned


def _decode_text(raw: bytes):
    for encoding in ("utf-8-sig", "utf-8", "cp1258"):
        try:
            return raw.decode(encoding), encoding, []
        except UnicodeDecodeError:
            continue
    return (raw.decode("latin-1", errors="replace"), "latin-1",
            ["Không xác định chắc chắn bảng mã; đã đọc dự phòng bằng latin-1."])


def _cell_text(value):
    if value is None:
        return ""
    if isinstance(value, (datetime, date, time)):
        return value.isoformat()
    text = re.sub(r"\s+", " ", str(value)).strip()
    return text[:MAX_CELL_CHARS]


def _validate_office_archive(path, invalid_code="invalid_office_archive"):
    try:
        with zipfile.ZipFile(path) as archive:
            expanded_size = sum(info.file_size for info in archive.infolist())
    except (OSError, zipfile.BadZipFile) as exc:
        raise ExtractionError(invalid_code, "File Office không phải gói ZIP hợp lệ.",
                              "Mở file bằng Office/Google Docs rồi lưu lại.") from exc
    if expanded_size > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
        raise ExtractionError(
            "office_archive_too_large",
            f"File Office sau giải nén vượt {MAX_ARCHIVE_UNCOMPRESSED_BYTES:,} byte.",
            "Tách file/bỏ ảnh nhúng quá lớn trước khi học.")


def _looks_texty(value: str) -> bool:
    """Ô 'chữ' — có ký tự chữ cái và không phải công thức. Dùng để phân biệt
    dòng tiêu đề cột với dòng tên báo cáo (ít ô) và dòng số liệu (toàn số)."""
    return bool(value) and not value.startswith("=") and any(ch.isalpha() for ch in value)


# Số dòng không rỗng đầu bảng được giữ lại để chọn dòng tiêu đề cột.
_HEADER_SCAN_ROWS = 8


def _pick_header(scanned):
    """Vị trí dòng TIÊU ĐỀ CỘT trong các dòng đầu bảng đã quét.

    Bản cũ lấy dòng không rỗng ĐẦU TIÊN làm tiêu đề. Báo cáo thực tế thường mở
    đầu bằng vài dòng tên công ty/tên báo cáo (1-2 ô có nội dung), nên tên cột
    thật bị đẩy xuống thành dữ liệu và mọi dòng sau mang tên cột vô nghĩa
    'Cột 3', 'Cột 8' — đoạn cắt ra tra cứu được mà không đọc được.

    Luật chọn: lấy dòng ĐẦU TIÊN có từ 2 ô trở lên, số ô không quá lép so với
    dòng dày nhất trong vùng quét (≥60%), và quá nửa số ô là chữ. Không dòng
    nào đạt thì lùi về dòng nhiều ô nhất, cuối cùng mới tới dòng đầu tiên
    (giữ hành vi cũ cho bảng một cột).
    """
    if not scanned:
        return 0
    max_count = max(sum(1 for v in values if v) for _, values in scanned)
    need = max(2, -(-max_count * 3 // 5))          # ceil(0.6 * max_count)
    for pos, (_, values) in enumerate(scanned):
        filled = [v for v in values if v]
        if len(filled) >= need and sum(_looks_texty(v) for v in filled) * 2 > len(filled):
            return pos
    for pos, (_, values) in enumerate(scanned):
        if sum(1 for v in values if v) == max_count:
            return pos
    return 0


def _tabular_text(rows, title, row_limit=MAX_SPREADSHEET_ROWS):
    """Biến bảng thành text có tên cột để truy vấn ngữ nghĩa chính xác hơn."""
    output = [f"[Bảng: {title}]"]
    headers = None
    data_rows = 0
    truncated = False
    cell_truncated = False
    column_truncated = False
    output_chars = len(output[0])
    scanned = []          # các dòng đầu, giữ tới khi chọn được dòng tiêu đề

    def emit(line):
        nonlocal output_chars, data_rows, truncated
        output.append(line)
        output_chars += len(line)
        data_rows += 1
        if output_chars >= MAX_EXTRACTED_CHARS:
            truncated = True
        return not truncated

    def build_headers(values):
        names, used = [], set()
        for index, value in enumerate(values, 1):
            base = value or f"Cột {index}"
            name = base
            duplicate = 2
            while name.casefold() in used:
                name = f"{base} ({duplicate})"
                duplicate += 1
            used.add(name.casefold())
            names.append(name)
        return names

    def data_line(source_index, values):
        if len(values) > len(headers):
            headers.extend(f"Cột {i}" for i in range(len(headers) + 1, len(values) + 1))
        cells = [f"{headers[i]}: {value}" for i, value in enumerate(values) if value]
        return f"[Dòng {source_index}] " + " | ".join(cells) if cells else None

    def flush_scanned():
        """Chọn tiêu đề trong các dòng đã quét rồi xả ra output đúng thứ tự."""
        nonlocal headers
        pending = scanned[:]
        scanned.clear()
        if not pending:
            return True
        pick = _pick_header(pending)
        for pos, (source_index, values) in enumerate(pending):
            if pos < pick:
                # Dòng mở đầu (tên báo cáo, kỳ báo cáo…) đứng TRƯỚC tiêu đề
                # cột: giữ nguyên nội dung, không ép vào khung 'tên cột: giá trị'.
                if not emit(f"[Dòng {source_index}] " + " | ".join(v for v in values if v)):
                    return False
            elif pos == pick:
                headers = build_headers(values)
                if not emit("[Cột] " + " | ".join(headers)):
                    return False
            else:
                line = data_line(source_index, values)
                if line and not emit(line):
                    return False
        return True

    for source_index, row in enumerate(rows, 1):
        if truncated:
            break
        if source_index > row_limit:
            truncated = True
            break
        row = list(row)
        if len(row) > MAX_SPREADSHEET_COLUMNS:
            row = row[:MAX_SPREADSHEET_COLUMNS]
            column_truncated = True
        values = [_cell_text(v) for v in row]
        if not any(values):
            continue
        if any(v is not None and len(str(v)) > MAX_CELL_CHARS for v in row):
            cell_truncated = True
        if headers is None:
            scanned.append((source_index, values))
            if len(scanned) >= _HEADER_SCAN_ROWS and not flush_scanned():
                break
            continue
        line = data_line(source_index, values)
        if line and not emit(line):
            break
    if headers is None:
        flush_scanned()

    warnings = []
    if truncated:
        warnings.append(
            f"Bảng đã bị giới hạn ở {row_limit:,} dòng hoặc {MAX_EXTRACTED_CHARS:,} ký tự; "
            "hãy tách file nếu cần tra cứu toàn bộ.")
    if cell_truncated:
        warnings.append(f"Một số ô dài quá {MAX_CELL_CHARS:,} ký tự đã được rút gọn.")
    if column_truncated:
        warnings.append(f"Bảng có quá {MAX_SPREADSHEET_COLUMNS:,} cột; các cột phía sau đã được bỏ qua.")
    return "\n".join(output if data_rows else []), data_rows, warnings, truncated


def _extract_txt(path):
    raw = path.read_bytes()
    text, encoding, warnings = _decode_text(raw)
    return text, "text", warnings, {"encoding": encoding}


def _extract_csv(path):
    raw = path.read_bytes()
    decoded, encoding, warnings = _decode_text(raw)
    sample = decoded[:16_384]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
        warnings.append("Không tự nhận diện được dấu phân cách CSV; đã dùng dấu phẩy.")
    try:
        text, rows, table_warnings, truncated = _tabular_text(
            csv.reader(io.StringIO(decoded), dialect), path.stem)
    except csv.Error as exc:
        raise ExtractionError("invalid_csv", "File CSV không hợp lệ.",
                              "Mở và lưu lại dưới dạng CSV UTF-8.") from exc
    warnings.extend(table_warnings)
    return text, "csv", warnings, {
        "encoding": encoding, "delimiter": dialect.delimiter, "rows": rows,
        "truncated": truncated,
    }


def _read_workbook(path, data_only):
    parts, warnings = [], []
    total_rows = 0
    visible_sheets = 0
    truncated = False
    from openpyxl import load_workbook
    workbook = load_workbook(str(path), read_only=True, data_only=data_only,
                             keep_links=False)
    try:
        for sheet in workbook.worksheets:
            if sheet.sheet_state != "visible":
                warnings.append(f"Đã bỏ qua sheet ẩn '{sheet.title}' để tránh học dữ liệu ngoài ý muốn.")
                continue
            visible_sheets += 1
            remaining = max(0, MAX_SPREADSHEET_ROWS - total_rows)
            if remaining == 0:
                truncated = True
                break
            text, rows, sheet_warnings, sheet_truncated = _tabular_text(
                sheet.iter_rows(values_only=True), sheet.title, remaining)
            if text:
                parts.append(text)
            total_rows += rows
            warnings.extend(sheet_warnings)
            truncated = truncated or sheet_truncated
            if sum(len(x) for x in parts) >= MAX_EXTRACTED_CHARS:
                truncated = True
                break
    finally:
        workbook.close()
    return parts, warnings, total_rows, visible_sheets, truncated


def _extract_xlsx(path):
    # data_only=True: đọc GIÁ TRỊ Excel đã tính sẵn, không phải công thức.
    # Bản cũ để data_only=False nên chunk chứa nguyên văn '=G14*8%', '=Q14*1.5%'
    # — model không tính được công thức, mọi câu hỏi số liệu đều thành vô dụng.
    # File do Excel/Google Sheets lưu luôn có sẵn giá trị cache; chỉ file sinh
    # bằng thư viện (chưa từng mở bằng Excel) mới thiếu — khi đó ô công thức trả
    # None và ta lùi về đọc công thức kèm cảnh báo, còn hơn mất trắng nội dung.
    try:
        _validate_office_archive(path, "invalid_xlsx")
        parts, warnings, total_rows, visible_sheets, truncated = _read_workbook(
            path, data_only=True)
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError("invalid_xlsx", "Không mở được file Excel.",
                              "Mở file bằng Excel/Google Sheets rồi lưu lại thành .xlsx.") from exc

    if not any(part.strip() for part in parts):
        try:
            parts, warnings, total_rows, visible_sheets, truncated = _read_workbook(
                path, data_only=False)
            if any(part.strip() for part in parts):
                warnings.append(
                    "File chưa lưu sẵn giá trị công thức nên phải đọc công thức thô; "
                    "mở file bằng Excel/Google Sheets rồi lưu lại để bot đọc được số liệu thật.")
        except Exception:
            pass

    if truncated and not any("giới hạn" in warning for warning in warnings):
        warnings.append(
            f"Workbook đã bị giới hạn ở {MAX_SPREADSHEET_ROWS:,} dòng hoặc "
            f"{MAX_EXTRACTED_CHARS:,} ký tự; hãy tách file nếu cần tra cứu toàn bộ.")
    return "\n\n".join(parts), "xlsx", warnings, {
        "sheets": visible_sheets, "rows": total_rows, "truncated": truncated,
    }


def _docx_xml_fallback(path):
    """Đọc text box khi python-docx không thấy nội dung; có chặn kích thước giải nén."""
    try:
        with zipfile.ZipFile(path) as archive:
            info = archive.getinfo("word/document.xml")
            if info.file_size > MAX_EXTRACTED_CHARS * 10:
                raise ExtractionError("docx_too_complex", "Nội dung DOCX sau giải nén quá lớn.",
                                      "Tách tài liệu thành các file nhỏ hơn.")
            root = ElementTree.fromstring(archive.read(info))
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError("invalid_docx", "Không mở được cấu trúc DOCX.",
                              "Mở file bằng Word/Google Docs rồi lưu lại thành .docx.") from exc
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = []
    for paragraph in root.findall(".//w:p", namespace):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", namespace)).strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)


def _extract_docx(path):
    try:
        _validate_office_archive(path, "invalid_docx")
        from docx import Document as Docx
        document = Docx(str(path))
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError("invalid_docx", "Không mở được file DOCX.",
                              "File có thể hỏng/đặt mật khẩu; hãy mở và lưu lại bằng Word.") from exc

    parts = []
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        if not paragraph_text:
            continue
        style = paragraph.style
        style_id = (getattr(style, "style_id", "") or "").casefold()
        style_name = (getattr(style, "name", "") or "").casefold()
        if (style_id.startswith("heading") or style_name.startswith("heading")
                or style_name.startswith("title") or style_name.startswith("tiêu đề")):
            parts.append(f"[Mục: {paragraph_text}]")
        else:
            parts.append(paragraph_text)
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            cells = [_cell_text(cell.text) for cell in row.cells]
            if any(cells):
                parts.append(f"[Bảng {table_index}, dòng {row_index}] " + " | ".join(cells))
    text = "\n".join(parts)
    warnings = []
    method = "python-docx"
    if len(clean(text)) < 20:
        fallback = _docx_xml_fallback(path)
        if len(clean(fallback)) > len(clean(text)):
            text = fallback
            method = "docx-xml"
            warnings.append("Đã dùng bộ đọc XML dự phòng vì nội dung nằm trong text box/đối tượng Word.")
    return text, method, warnings, {"tables": len(document.tables)}


def _extract_pdf(path):
    pages = []
    pdf_error = None
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            for page_number, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(f"[Trang {page_number}]\n{page_text}")
            page_count = len(pdf.pages)
    except Exception as exc:
        pdf_error = exc
        page_count = 0

    text = "\n\n".join(pages)
    warnings = []
    method = "pdfplumber"
    if len(clean(text)) < 100:
        try:
            ocr_text = _ocr_pdf_strict(path)
        except ExtractionError as exc:
            if not clean(text):
                if pdf_error:
                    raise ExtractionError("unreadable_pdf", "Không đọc được PDF và OCR thất bại.",
                                          exc.hint) from pdf_error
                raise
            warnings.append(f"Nội dung PDF rất ít và OCR không chạy được: {exc.hint or exc.message}")
        else:
            if clean(ocr_text):
                text = ocr_text
                method = "ocr"
                warnings.append("PDF dạng scan/ít text; đã dùng OCR tiếng Việt.")
    return text, method, warnings, {"pages": page_count, "pdf_error": bool(pdf_error)}


def extract_text_with_metadata(path: Path) -> ExtractionResult:
    """Trích nội dung và trả method/warnings/metadata cho status vận hành."""
    path = Path(path)
    try:
        size = path.stat().st_size
    except OSError as exc:
        raise ExtractionError("file_unreadable", "Không đọc được file nguồn.",
                              "Kiểm tra file còn tồn tại và quyền đọc.") from exc
    if not path.is_file():
        raise ExtractionError("not_a_file", "Đường dẫn nguồn không phải là file.")
    if size == 0:
        raise ExtractionError("empty_file", "File rỗng, không có nội dung để học.")
    if size > MAX_SOURCE_BYTES:
        raise ExtractionError("file_too_large", f"File lớn hơn giới hạn {MAX_SOURCE_BYTES:,} byte.",
                              "Tách file thành các phần nhỏ hơn hoặc tăng INGEST_MAX_SOURCE_BYTES có kiểm soát.")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ExtractionError("unsupported_format", f"Chưa hỗ trợ định dạng '{ext or '(không có đuôi)'}'.",
                              "Dùng PDF, DOCX, DOC, TXT, MD, XLSX hoặc CSV.")
    try:
        if ext in (".txt", ".md"):
            text, method, warnings, metadata = _extract_txt(path)
        elif ext == ".csv":
            text, method, warnings, metadata = _extract_csv(path)
        elif ext == ".xlsx":
            text, method, warnings, metadata = _extract_xlsx(path)
        elif ext == ".docx":
            text, method, warnings, metadata = _extract_docx(path)
        elif ext == ".doc":
            text = _extract_doc_strict(path)
            method, warnings, metadata = "libreoffice", [], {}
        else:
            text, method, warnings, metadata = _extract_pdf(path)
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError("extraction_failed", "Trích xuất nội dung thất bại.",
                              f"Kiểm tra file {path.name} có bị hỏng hoặc đặt mật khẩu.") from exc

    text = clean(text)
    if not text or not any(ch.isalnum() for ch in text):
        raise ExtractionError("no_text", "Không tìm thấy nội dung chữ có thể dùng.",
                              "Nếu tài liệu là ảnh/scan, kiểm tra OCR; nếu là DOCX, mở và lưu lại file.")
    if len(text) < 80:
        warnings.append("Nội dung trích xuất dưới 80 ký tự; cần kiểm tra thủ công trước khi duyệt.")
    metadata = {**metadata, "source_bytes": size, "characters": len(text)}
    return ExtractionResult(text=text, format=ext.lstrip("."), method=method,
                            warnings=warnings, metadata=metadata)


def extract_text(path: Path) -> str:
    """API tương thích cũ: trả chuỗi rỗng khi lỗi và in lý do có mã."""
    try:
        return extract_text_with_metadata(path).text
    except ExtractionError as exc:
        print(f"  [!] [{exc.code}] {exc}")
        return ""


def _extract_doc_strict(path: Path) -> str:
    """Đọc file .doc (Word 97-2003) bằng cách nhờ LibreOffice chuyển sang .docx
    trong thư mục tạm rồi đọc lại. Rất nhiều tài liệu cũ của HDS ở định dạng
    này; không đọc được chúng thì bot thiếu hẳn một mảng dữ liệu lớn.

    Cần gói hệ thống 'libreoffice' trên máy chủ (deploy/setup.sh đã cài).
    """
    import shutil
    import subprocess
    import tempfile

    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        raise ExtractionError("libreoffice_missing", "Chưa có LibreOffice để đọc file .doc.",
                              "Cài libreoffice hoặc chuyển file sang .docx.")
    try:
        with tempfile.TemporaryDirectory() as tmp:
            # UserInstallation riêng để nhiều lần gọi liên tiếp không khoá hồ sơ nhau
            subprocess.run(
                [soffice, f"-env:UserInstallation=file://{tmp}/profile",
                 "--headless", "--convert-to", "docx", "--outdir", tmp, str(path)],
                check=True, timeout=120,
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            out = Path(tmp) / (path.stem + ".docx")
            if not out.exists():
                raise ExtractionError("doc_conversion_failed", "LibreOffice không tạo được file DOCX.",
                                      "Mở file .doc và lưu lại thủ công thành .docx.")
            return _extract_docx(out)[0]
    except ExtractionError:
        raise
    except Exception as e:
        raise ExtractionError("doc_conversion_failed", "Không chuyển được .doc sang .docx.",
                              "File có thể hỏng/đặt mật khẩu; hãy lưu lại thành .docx.") from e


def extract_doc(path: Path) -> str:
    """API tương thích cũ cho bộ đọc .doc."""
    try:
        return _extract_doc_strict(path)
    except ExtractionError as exc:
        print(f"  [!] [{exc.code}] {path.name}: {exc}")
        return ""


def _ocr_pdf_strict(path: Path) -> str:
    try:
        import pytesseract
        from pdf2image import convert_from_path, pdfinfo_from_path
        page_count = int(pdfinfo_from_path(str(path)).get("Pages") or 0)
        if page_count > MAX_OCR_PAGES:
            raise ExtractionError(
                "ocr_page_limit", f"PDF scan có {page_count} trang, vượt giới hạn OCR {MAX_OCR_PAGES} trang.",
                "Tách PDF thành các phần nhỏ hơn hoặc tăng INGEST_MAX_OCR_PAGES có kiểm soát.")
        parts = []
        # OCR theo lô nhỏ để không giữ ảnh của cả PDF trong RAM cùng lúc.
        for first_page in range(1, page_count + 1, 10):
            last_page = min(page_count, first_page + 9)
            images = convert_from_path(str(path), dpi=300, first_page=first_page, last_page=last_page)
            for offset, image in enumerate(images):
                page_number = first_page + offset
                value = pytesseract.image_to_string(image, lang="vie")
                if value.strip():
                    parts.append(f"[Trang {page_number}]\n{value}")
        return "\n\n".join(parts)
    except ExtractionError:
        raise
    except Exception as e:
        raise ExtractionError("ocr_failed", "OCR PDF thất bại.",
                              "Cài tesseract-ocr-vie và poppler-utils, hoặc dùng PDF có lớp text.") from e


def ocr_pdf(path: Path) -> str:
    """API tương thích cũ cho OCR."""
    try:
        return _ocr_pdf_strict(path)
    except ExtractionError as exc:
        print(f"  [!] [{exc.code}] {exc}")
        return ""


def clean(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


RE_DIEU = re.compile(r"^\s*(Điều\s+\d+[a-z]?)\s*[.:]?", re.MULTILINE | re.IGNORECASE)


def document_citation(text: str) -> str:
    """Đọc SỐ HIỆU VĂN BẢN ở đầu văn bản luật, ví dụ 'Thông tư 01/2021/TT-BXD'.

    Không có dòng này thì mọi đoạn cắt ra đều là 'Điều 5' trơ trọi — người đọc
    không biết Điều 5 của văn bản nào, và câu trả lời của bot không dẫn nguồn
    được theo chuẩn hành nghề. Vì vậy số hiệu phải được gắn vào TỪNG đoạn.

    Chỉ soi phần đầu: số hiệu luôn nằm ở phần mở đầu, còn thân văn bản thì đầy
    số hiệu của văn bản KHÁC (điều khoản dẫn chiếu) — quét cả bài là lấy nhầm.
    """
    head = text[:4000]
    # 45/2019/QH14, 01/2021/TT-BXD, 15/2020/NĐ-CP…
    number = re.search(
        r"\bS[ốô]\s*:?\s*(\d+\s*/\s*\d{4}\s*/\s*[A-ZĐ][A-ZĐ0-9\-]*)", head)
    if not number:
        number = re.search(r"\b(\d+/\d{4}/[A-ZĐ][A-ZĐ0-9\-]+)\b", head)
    kind = re.search(
        r"\b(BỘ LUẬT|LUẬT|PHÁP LỆNH|NGHỊ ĐỊNH|NGHỊ QUYẾT|THÔNG TƯ LIÊN TỊCH|"
        r"THÔNG TƯ|QUYẾT ĐỊNH|CHỈ THỊ|CÔNG VĂN|Bộ luật|Luật|Pháp lệnh|"
        r"Nghị định|Nghị quyết|Thông tư liên tịch|Thông tư|Quyết định)\b", head)
    parts = []
    if kind:
        parts.append(kind.group(1).title() if kind.group(1).isupper() else kind.group(1))
    if number:
        parts.append("số " + re.sub(r"\s+", "", number.group(1)))
    return " ".join(parts).strip()


# Cấp bậc trong văn bản quy phạm pháp luật Việt Nam, từ lớn tới nhỏ.
RE_PHAN = re.compile(r"^\s*(Phần\s+(?:thứ\s+)?[^\n]{0,60})$", re.MULTILINE | re.IGNORECASE)
RE_CHUONG = re.compile(r"^\s*(Chương\s+[IVXLCDM\d]+[^\n]{0,80})$", re.MULTILINE | re.IGNORECASE)
RE_MUC = re.compile(r"^\s*(Mục\s+\d+[^\n]{0,80})$", re.MULTILINE | re.IGNORECASE)


def _heading_path(text: str, position: int) -> str:
    """Chương/Mục đang có hiệu lực tại vị trí này — tiêu đề gần nhất phía trên.

    Điều 5 nằm trong Chương II khác hẳn Điều 5 của Chương V. Không giữ đường dẫn
    này thì hai điều trùng số bị trộn vào nhau khi tra cứu.
    """
    path = []
    for pattern in (RE_PHAN, RE_CHUONG, RE_MUC):
        last = None
        for match in pattern.finditer(text):
            if match.start() < position:
                last = match.group(1).strip()
            else:
                break
        if last:
            path.append(re.sub(r"\s+", " ", last))
    return ", ".join(path)


RE_TRANG = re.compile(r"^\[Trang\s+(\d+)\]\s*$", re.MULTILINE | re.IGNORECASE)


def _page_at(text: str, position: int):
    """Số trang PDF của vị trí này — mốc [Trang n] gần nhất phía trên.

    Cắt luật theo điều nên không còn đi theo từng trang nữa, nhưng số trang vẫn
    cần cho người muốn mở đúng chỗ trong file gốc.
    """
    page = None
    for match in RE_TRANG.finditer(text):
        if match.start() < position:
            page = int(match.group(1))
        else:
            break
    return page


def chunk_law_structured(text: str) -> list[ChunkPiece]:
    """Cắt văn bản luật theo ĐIỀU, giữ nguyên đường dẫn trích dẫn của từng điều.

    Mỗi đoạn trả về đều tự mang đủ thông tin để trích dẫn: số hiệu văn bản,
    Chương/Mục, và số Điều. Nhờ vậy khi bot dẫn nguồn, người đọc tra ngược được
    tới đúng điều khoản mà không phải mở lại cả file.

    Điều quá dài vẫn phải cắt nhỏ, nhưng mỗi phần đều lặp lại nhãn để không có
    mảnh nào mất danh tính.
    """
    citation = document_citation(text)
    marks = [(m.start(), m.group(1).strip()) for m in RE_DIEU.finditer(text)]
    if not marks:
        return [ChunkPiece(content=piece,
                           section_title=citation or None,
                           source_locator="document")
                for piece in chunk_generic(text)]

    pieces: list[ChunkPiece] = []
    for index, (position, label) in enumerate(marks):
        end = marks[index + 1][0] if index + 1 < len(marks) else len(text)
        body = text[position:end].strip()
        if len(body) < 20:
            continue
        path = _heading_path(text, position)
        page = _page_at(text, position)
        # Tiêu đề đầy đủ: "Thông tư số 01/2021/TT-BXD — Chương II, Mục 1, Điều 5"
        section = " — ".join(x for x in (citation, path, label) if x)
        number = re.search(r"(\d+[a-z]?)", label)
        locator = f"dieu:{number.group(1)}" if number else "document"

        if len(body.split()) > CHUNK_WORDS * 2:
            parts = chunk_generic(body)
            for order, sub in enumerate(parts, 1):
                # Nhãn nằm TRONG nội dung, không chỉ ở metadata: đoạn được đưa
                # vào prompt dưới dạng văn bản thuần, model chỉ đọc được cái gì
                # nằm trong nội dung.
                header = f"[{section} — phần {order}/{len(parts)}]"
                pieces.append(ChunkPiece(
                    content=f"{header}\n{sub}", page_number=page,
                    section_title=section,
                    source_locator=f"{locator};phan:{order}"))
        else:
            header = f"[{section}]" if citation or path else ""
            content = f"{header}\n{body}" if header else body
            pieces.append(ChunkPiece(content=content, page_number=page,
                                     section_title=section,
                                     source_locator=locator))
    return pieces


def chunk_law(text):
    """Bản trả về chuỗi, giữ cho các chỗ gọi cũ (split_document) không vỡ."""
    return [piece.content for piece in chunk_law_structured(text)]


# Viết tắt tiếng Việt hay gặp — dấu chấm sau chúng KHÔNG kết thúc câu.
# Thiếu danh sách này thì "Nghị định số 01/2021/NĐ-CP." hay "TP. Hà Nội" bị
# tách làm đôi, và đoạn cắt ra mất nghĩa ngay giữa một cụm danh từ.
_ABBREV = {
    "tp", "tt", "nđ", "cp", "qh", "ubnd", "hđnd", "tnhh", "vd", "vv", "stt",
    "ts", "ths", "gs", "pgs", "bs", "ls", "kts", "ks", "nxb", "tr", "gđ",
    "ông", "bà", "số", "kèm", "đ", "đv", "tk", "hđ", "hđlđ",
}

# Ranh giới câu: dấu kết câu + khoảng trắng + chữ hoa/số bắt đầu câu mới.
_SENTENCE_SPLIT = re.compile(r"(?<=[.!?…])\s+(?=[\"'“(\[]?[A-ZĐÀ-Ỹ0-9])")


def _is_heading(line: str) -> bool:
    """Dòng này có phải TIÊU ĐỀ MỤC không — chỗ được phép ngắt đoạn.

    Ngắt ở tiêu đề luôn tốt hơn ngắt giữa chừng: mỗi đoạn ra sẽ trọn một ý,
    và tiêu đề đi kèm cho biết đoạn đó nói về cái gì.
    """
    s = (line or "").strip()
    if not s or len(s) > 120:
        return False
    if re.match(r"^#{1,6}\s", s):                       # markdown
        return True
    if re.match(r"^(Điều|Chương|Mục|Phần|Khoản)\s+[\dIVXLC]", s, re.I):
        return True
    if re.match(r"^([IVXLC]+|\d+(\.\d+)*|[A-Z])[.)]\s+\S", s):   # 1. / 1.1. / I. / A.
        return True
    if re.match(r"^\[(Bảng|Trang|Cột)\b", s, re.I):     # mốc do bộ trích xuất chèn
        return True
    letters = [c for c in s if c.isalpha()]
    if len(letters) >= 3 and all(c.isupper() for c in letters):
        return True                                     # DÒNG VIẾT HOA TOÀN BỘ
    return s.endswith(":") and len(s) <= 80


def _split_units(text: str) -> list[str]:
    """Chia văn bản thành ĐƠN VỊ KHÔNG ĐƯỢC PHÉP CẮT ĐÔI.

    Mỗi đơn vị là một câu trọn vẹn, hoặc một dòng với nội dung theo dòng (dòng
    bảng tính `[Dòng 5] …`, gạch đầu dòng). Cắt giữa câu là cách chắc chắn nhất
    để một đoạn trở nên vô nghĩa với cả người đọc lẫn model.
    """
    units: list[str] = []
    for para in re.split(r"\n\s*\n|\n", text or ""):
        para = para.strip()
        if not para:
            continue
        # Nội dung theo dòng (bảng, danh sách, tiêu đề) giữ nguyên cả dòng.
        if _is_heading(para) or re.match(r"^[-·•*]\s", para):
            units.append(para)
            continue
        parts = _SENTENCE_SPLIT.split(para)
        buffer = ""
        for part in parts:
            candidate = (buffer + " " + part).strip() if buffer else part
            # Kết thúc bằng viết tắt thì câu chưa hết — nối tiếp phần sau.
            last = re.sub(r"[^\wÀ-ỹ]", "", candidate.split()[-1]).lower() if candidate.split() else ""
            if last in _ABBREV:
                buffer = candidate
                continue
            units.append(candidate)
            buffer = ""
        if buffer:
            units.append(buffer)
    return units


def _cohesion(left: list[str], right: list[str]) -> float:
    """Mạch văn giữa hai bên có liền nhau không — đo bằng từ dùng chung.

    Ý tưởng của TextTiling: hai phần cùng bàn một chuyện thì lặp lại nhiều từ;
    khi chủ đề chuyển, lượng từ chung tụt xuống. Chỗ tụt sâu nhất chính là chỗ
    nên ngắt. Rẻ (chỉ đếm từ), không cần gọi model, và cho kết quả y hệt nhau ở
    mọi lần chạy — điều kiện cần để lỗi tra cứu còn tái hiện được.
    """
    lt = {w for w in re.findall(r"[\wÀ-ỹ]+", " ".join(left).lower()) if len(w) > 2}
    rt = {w for w in re.findall(r"[\wÀ-ỹ]+", " ".join(right).lower()) if len(w) > 2}
    if not lt or not rt:
        return 0.0
    return len(lt & rt) / min(len(lt), len(rt))


def chunk_generic(text, target_words=None, overlap_units=CHUNK_OVERLAP_UNITS):
    """Cắt đoạn THEO NGỮ CẢNH: tiêu đề → câu trọn vẹn → chỗ mạch văn đứt.

    Bản cũ cắt cứng mỗi 320 từ bất kể nội dung, nên thường xuyên cắt ngang câu,
    ngang điều khoản, ngang một dòng bảng. Đoạn hỏng kiểu đó vẫn được tạo vector
    và vẫn được trả về khi tra cứu — model đọc phải một mẩu cụt và trả lời cụt
    theo.

    Ba tầng, từ chắc chắn nhất tới suy đoán nhất:
      1. Ngắt ở TIÊU ĐỀ MỤC — ranh giới do chính tác giả văn bản đặt ra.
      2. Trong mỗi mục, gom CÂU TRỌN VẸN cho tới khi đủ dài; không bao giờ cắt
         giữa câu.
      3. Mục dài hơn ngân sách thì chọn điểm ngắt ở chỗ hai bên ÍT DÙNG CHUNG
         TỪ NHẤT trong vùng cho phép — tức chỗ chủ đề chuyển.

    Kích thước đoạn vì vậy KHÔNG cố định: mục ngắn giữ nguyên một đoạn, mục dài
    tự tách ở đúng chỗ chuyển ý.
    """
    target = target_words or CHUNK_WORDS
    units = _split_units(text)
    if not units:
        return []

    def wc(items):
        return sum(len(u.split()) for u in items)

    # ---- Tầng 1: gom theo mục, mỗi tiêu đề mở một mục mới ----------------
    sections: list[list[str]] = []
    current: list[str] = []
    for unit in units:
        if _is_heading(unit) and current:
            sections.append(current)
            current = [unit]
        else:
            current.append(unit)
    if current:
        sections.append(current)

    # Mục quá ngắn thì nhập vào mục kế: một tiêu đề trơ trọi không thành đoạn.
    merged: list[list[str]] = []
    for section in sections:
        if merged and wc(merged[-1]) < target * 0.35:
            merged[-1].extend(section)
        else:
            merged.append(section)
    sections = merged

    # ---- Tầng 2+3: trong mỗi mục, gom câu tới hạn rồi tìm chỗ ngắt tốt ----
    out: list[str] = []
    for section in sections:
        if wc(section) <= target:
            out.append(" ".join(section).strip())      # đủ ngắn → giữ trọn mục
            continue
        start = 0
        while start < len(section):
            end, count = start, 0
            while end < len(section) and count < target:
                count += len(section[end].split())
                end += 1
            # Vùng được phép xê dịch điểm ngắt: 25% cuối đoạn. Trong vùng đó
            # chọn chỗ hai bên ít dùng chung từ nhất — chỗ chuyển ý rõ nhất.
            if end < len(section):
                low = max(start + 1, end - max(1, (end - start) // 4))
                best, best_score = end, 2.0
                for cut in range(low, end + 1):
                    score = _cohesion(section[start:cut], section[cut:cut + 3])
                    if score < best_score:
                        best, best_score = cut, score
                end = best
            piece = " ".join(section[start:end]).strip()
            if piece:
                out.append(piece)
            if end >= len(section):
                break
            # Chồng lấn bằng CẢ CÂU, không phải vài chục từ lẻ: câu đầu đoạn sau
            # nhắc lại ý cuối đoạn trước để không mất mạch khi đọc rời từng đoạn.
            start = max(start + 1, end - overlap_units)
    return [piece for piece in out if piece]


def split_document(text, doc_type):
    text = clean(text)
    if not text:
        return []
    return chunk_law(text) if doc_type == "law" else chunk_generic(text)


def _segments_from_extraction(extraction: ExtractionResult):
    text = extraction.text
    if extraction.format == "pdf":
        pattern = re.compile(r"^\[Trang\s+(\d+)\]\s*$", re.MULTILINE | re.IGNORECASE)
        matches = list(pattern.finditer(text))
        if matches:
            segments = []
            for index, match in enumerate(matches):
                end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
                page_number = int(match.group(1))
                content = text[match.start():end].strip()
                if content:
                    segments.append(_SourceSegment(
                        content=content, page_number=page_number,
                        source_locator=f"page:{page_number}"))
            return segments

    if extraction.format in {"xlsx", "csv"}:
        pattern = re.compile(r"^\[Bảng:\s*(.+?)\]\s*$", re.MULTILINE | re.IGNORECASE)
        matches = list(pattern.finditer(text))
        if matches:
            segments = []
            for index, match in enumerate(matches):
                end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
                title = match.group(1).strip()
                content = text[match.start():end].strip()
                prefix = "sheet" if extraction.format == "xlsx" else "table"
                if content:
                    segments.append(_SourceSegment(
                        content=content, section_title=title,
                        source_locator=f"{prefix}:{title}"))
            return segments

    if extraction.format in {"docx", "doc"}:
        pattern = re.compile(r"^\[Mục:\s*(.+?)\]\s*$", re.MULTILINE | re.IGNORECASE)
        matches = list(pattern.finditer(text))
        if matches:
            segments = []
            if text[:matches[0].start()].strip():
                segments.append(_SourceSegment(content=text[:matches[0].start()].strip()))
            for index, match in enumerate(matches):
                end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
                title = match.group(1).strip()
                content = text[match.start():end].strip()
                if content:
                    segments.append(_SourceSegment(
                        content=content, section_title=title,
                        source_locator=f"section:{title}"))
            return segments

    return [_SourceSegment(content=text)]


def _chunk_locator(segment: _SourceSegment, content: str, source_format: str):
    if source_format not in {"xlsx", "csv"}:
        return segment.source_locator
    rows = [int(value) for value in re.findall(r"\[Dòng\s+(\d+)\]", content,
                                               flags=re.IGNORECASE)]
    if not rows:
        return f"{segment.source_locator};header"
    return f"{segment.source_locator};rows:{min(rows)}-{max(rows)}"


def context_header(title: str, doc_type: str, client_name: str | None = None) -> str:
    """Dòng DANH TÍNH gắn vào đầu mỗi đoạn TRƯỚC khi tạo vector.

    Đoạn cắt rời khỏi file không tự biết nó là của ai: 'tổng số lao động: 02'
    trong hồ sơ đăng ký doanh nghiệp CỦA KHÁCH đọc y hệt số liệu của chính HDS.
    Nhúng danh tính vào nội dung được embedding có hai tác dụng: câu hỏi về
    'công ty tôi' bớt khớp nhầm với hồ sơ khách, và khi đoạn được trích dẫn thì
    cả model lẫn người đọc thấy ngay nó thuộc hồ sơ nào.
    """
    try:
        from app.company_context import DOC_TYPE_VN
    except Exception:
        DOC_TYPE_VN = {}
    bits = [f"Tài liệu: {title}"] if title else []
    if client_name:
        bits.append(f"hồ sơ khách hàng — {client_name}")
    elif doc_type == "ho_so_ns":
        bits.append("hồ sơ nhân sự của công ty luật HDS")
    else:
        label = DOC_TYPE_VN.get(doc_type or "")
        if label and doc_type != "other":
            bits.append(label)
    return f"[{' | '.join(bits)}]" if bits else ""


def apply_context_headers(pieces, title, doc_type, client_name=None):
    """Gắn dòng danh tính vào từng đoạn. Đoạn LUẬT giữ nguyên — chúng đã tự
    mang số hiệu văn bản + Chương/Điều trong nội dung (chunk_law_structured)."""
    if doc_type == "law":
        return pieces
    header = context_header(title, doc_type, client_name)
    if not header:
        return pieces
    for piece in pieces:
        piece.content = f"{header}\n{piece.content}"
    return pieces


def client_display_name(client_id):
    """Tên khách cho dòng danh tính; None khi tài liệu không thuộc khách nào."""
    if not client_id:
        return None
    try:
        with db.session(role="internal", admin=True) as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT name FROM clients WHERE id=%s", (client_id,))
                row = cur.fetchone()
        return row[0] if row else None
    except Exception:
        return None


def split_document_with_metadata(extraction: ExtractionResult, doc_type):
    """Chia đoạn nhưng không làm mất trang/mục/sheet dùng cho trích dẫn chính xác."""
    output = []
    # Văn bản luật phải được cắt trên TOÀN VĂN, không cắt theo từng trang PDF:
    # một Điều thường vắt qua hai trang, và Chương/Mục thì nằm cách đó vài chục
    # trang. Cắt theo trang là băm nát điều khoản và mất luôn đường dẫn trích dẫn.
    if doc_type == "law":
        full_text = clean(extraction.text)
        if full_text:
            for piece in chunk_law_structured(full_text):
                output.append(piece)
        return output

    for segment in _segments_from_extraction(extraction):
        segment_text = clean(segment.content)
        if not segment_text:
            continue
        for piece in chunk_generic(segment_text):
            output.append(ChunkPiece(
                content=piece,
                page_number=segment.page_number,
                section_title=segment.section_title,
                source_locator=_chunk_locator(segment, piece, extraction.format),
            ))
    return output


def ingest_file(path: Path, doc_type="other", access_level="internal", client_id=None,
                department_id=None, matter_id=None,
                approved=False, label_verified=False, source_kind="manual"):
    print(f">> {path.name}")
    if access_level == "client" and client_id is None:
        print("  [BỎ QUA] client thiếu client_id → nguy cơ lộ dữ liệu chéo.")
        return None
    try:
        extraction = extract_text_with_metadata(path)
    except ExtractionError as exc:
        print(f"  [BỎ QUA] [{exc.code}] {exc}")
        return None
    text = extraction.text
    print(f"  Trích xuất bằng {extraction.method}; {len(text):,} ký tự.")
    for warning in extraction.warnings:
        print(f"  [CẢNH BÁO] {warning}")
    pieces = split_document_with_metadata(extraction, doc_type)
    if not pieces:
        print("  [BỎ QUA] không chia được đoạn.")
        return None
    pieces = apply_context_headers(pieces, path.stem, doc_type,
                                   client_display_name(client_id))
    print(f"  {len(text)} ký tự → {len(pieces)} đoạn, đang tạo vector...")
    vecs = embed([piece.content for piece in pieces])
    print("  Đang tạo tóm tắt...")
    summary = summarize(text, path.stem)
    checksum = hashlib.md5(text.encode()).hexdigest()
    extraction_status = "warning" if extraction.warnings else "ready"
    extraction_note = (json.dumps(extraction.warnings, ensure_ascii=False)
                       if extraction.warnings else None)
    # Cảnh báo extraction luôn cần người duyệt; không để `auto_approve=true`
    # biến một OCR/bảng bị cắt thành nguồn sự thật ngay lập tức.
    safe_approved = bool(approved and extraction_status == "ready")
    safe_verified = bool(label_verified and extraction_status == "ready")
    with db.session(role="internal", admin=True) as conn:
        with conn.cursor() as cur:
            cur.execute("""INSERT INTO documents
                (title,source_path,checksum,doc_type,access_level,client_id,department_id,matter_id,
                 approved,label_verified,source_kind,summary,extraction_status,extraction_error)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id""",
                (path.stem, str(path), checksum, doc_type, access_level, client_id, department_id, matter_id,
                 safe_approved, safe_verified, source_kind, summary,
                 extraction_status, extraction_note))
            doc_id = cur.fetchone()[0]
            for idx, (piece, vec) in enumerate(zip(pieces, vecs)):
                cur.execute("""INSERT INTO chunks
                    (document_id,chunk_index,content,page_number,section_title,source_locator,
                     access_level,client_id,department_id,doc_type,embedding)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                    (doc_id, idx, piece.content, piece.page_number, piece.section_title,
                     piece.source_locator, access_level, client_id, department_id,
                     doc_type, json.dumps(vec)))
        db.audit(conn, None, "ingest_document", "documents", doc_id,
                 {"file": path.name, "chunks": len(pieces),
                  "extraction": {"method": extraction.method,
                                 "status": extraction.status,
                                 "warnings": extraction.warnings,
                                 "metadata": extraction.metadata}})
    print(f"  [OK] document_id={doc_id}, {len(pieces)} đoạn.")
    return doc_id


def ingest_folder(folder, **kw):
    root = Path(folder)
    if not root.exists():
        print(f"Không thấy thư mục: {folder}")
        return
    files = [p for p in root.rglob("*") if p.suffix.lower() in SUPPORTED_EXTENSIONS]
    print(f"Tìm thấy {len(files)} file.\n")
    ok = 0
    for p in files:
        try:
            if ingest_file(p, **kw):
                ok += 1
        except Exception as e:
            print(f"  [LỖI] {p.name}: {e}")
    print(f"\nHoàn tất: {ok}/{len(files)} file.")


if __name__ == "__main__":
    folder = sys.argv[1] if len(sys.argv) > 1 else "data/raw"
    dtype = sys.argv[2] if len(sys.argv) > 2 else "other"
    ingest_folder(folder, doc_type=dtype, access_level="internal",
                  approved=True, label_verified=True)
