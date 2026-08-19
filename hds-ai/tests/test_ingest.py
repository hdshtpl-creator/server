"""Unit test thuần cho trích xuất; không cần PostgreSQL, Ollama hay Google Drive.

Chạy: python -m unittest tests.test_ingest -v
"""
import io
import tempfile
import unittest
from contextlib import redirect_stdout
from pathlib import Path
from unittest.mock import patch

from app import auto_learn
from app.auto_learn import auto_approve_from_env, drive_fingerprint
from app.ingest import (ExtractionError, ExtractionResult, _needs_ocr,
                        _pdf_text_via_pypdf, _split_units,
                        chunk_generic, chunk_law_structured,
                        document_citation, extract_text,
                        extract_text_with_metadata, safe_path_component,
                        split_document_with_metadata)


class PdfResilienceTests(unittest.TestCase):
    """Ca thật 19/08/2026: giấy tờ scan (PDF) của một nhân sự lỗi trích xuất
    và bị BỎ QUA — người đó biến mất khỏi mọi câu trả lời nhân sự. PDF hỏng
    một tầng không được phép chết cả đường đọc."""

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)

    def tearDown(self):
        self.temp_dir.cleanup()

    # ---- _needs_ocr: quyết theo MẬT ĐỘ chữ mỗi trang ----
    def test_needs_ocr_when_text_tiny(self):
        self.assertTrue(_needs_ocr("chỉ vài chữ", 1))

    def test_needs_ocr_scan_with_header_text(self):
        # Scan 3 trang nhưng chỉ có ~180 ký tự header thật — ngưỡng cũ (<100
        # tổng) bỏ sót đúng loại này, tài liệu vào kho chỉ với mấy dòng header.
        header = "CÔNG TY LUẬT HDS — HỒ SƠ NHÂN SỰ NỘI BỘ. " * 4
        self.assertTrue(_needs_ocr(header, 3))

    def test_no_ocr_for_real_text_pdf(self):
        page = "Điều 1. Nội dung có chữ đầy đủ, mỗi trang vài trăm ký tự. " * 10
        self.assertFalse(_needs_ocr(page * 2, 2))

    def test_no_ocr_when_pages_unknown_but_text_rich(self):
        page = "Nội dung dài và có nghĩa. " * 20
        self.assertFalse(_needs_ocr(page, 0))

    # ---- pypdf: tầng cứu khi pdfplumber bó tay ----
    def test_pypdf_reads_page_count_of_valid_pdf(self):
        from pypdf import PdfWriter
        path = self.root / "blank.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=595, height=842)
        with open(path, "wb") as fh:
            writer.write(fh)
        text, pages = _pdf_text_via_pypdf(path)
        self.assertEqual(pages, 1)          # đường pypdf sống, đếm đúng trang
        self.assertEqual(text, "")          # trang trắng thì không có chữ

    def test_pypdf_returns_empty_on_garbage(self):
        path = self.root / "rac.pdf"
        path.write_bytes(b"day khong phai pdf" * 20)
        self.assertEqual(_pdf_text_via_pypdf(path), ("", 0))

    # ---- Toàn tuyến: file rác phải ra ExtractionError có mã, không crash lạ ----
    def test_garbage_pdf_fails_with_coded_error(self):
        path = self.root / "hong.pdf"
        path.write_bytes(b"%PDF-1.4 nhung ruot thi hong het" + b"\x00" * 200)
        with self.assertRaises(ExtractionError) as ctx:
            extract_text_with_metadata(path)
        self.assertEqual(ctx.exception.code, "unreadable_pdf")
        # Hint phải nói được bước nào hỏng để dashboard chỉ đúng cách sửa.
        self.assertTrue(ctx.exception.hint)


class IngestExtractionTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_csv_keeps_headers_and_vietnamese_values(self):
        path = self.root / "nhan-su.csv"
        path.write_text(
            "Họ tên;Tình trạng;Ngày hết hạn\n"
            "Nguyễn Văn An;Còn hiệu lực;2027-01-10\n",
            encoding="utf-8",
        )

        result = extract_text_with_metadata(path)

        self.assertEqual("csv", result.format)
        self.assertEqual(";", result.metadata["delimiter"])
        self.assertIn("Họ tên: Nguyễn Văn An", result.text)
        self.assertIn("Tình trạng: Còn hiệu lực", result.text)

    def test_xlsx_keeps_sheet_and_columns_but_skips_hidden_sheet(self):
        from openpyxl import Workbook

        path = self.root / "hop-dong.xlsx"
        workbook = Workbook()
        visible = workbook.active
        visible.title = "Hợp đồng"
        visible.append(["Nhân sự", "Ngày hết hạn"])
        visible.append(["Trần Bình", "2028-12-31"])
        hidden = workbook.create_sheet("Dữ liệu ẩn")
        hidden.append(["Mật khẩu", "không được học"])
        hidden.sheet_state = "hidden"
        workbook.save(path)

        result = extract_text_with_metadata(path)

        self.assertEqual("xlsx", result.method)
        self.assertIn("[Bảng: Hợp đồng]", result.text)
        self.assertIn("Nhân sự: Trần Bình", result.text)
        self.assertNotIn("không được học", result.text)
        self.assertTrue(any("sheet ẩn" in warning for warning in result.warnings))
        chunks = split_document_with_metadata(result, "other")
        self.assertEqual("Hợp đồng", chunks[0].section_title)
        self.assertIn("sheet:Hợp đồng;rows:2-2", chunks[0].source_locator)

    def test_xlsx_title_rows_do_not_become_column_headers(self):
        """Lỗi thực tế 19/08/2026: báo cáo mở đầu bằng dòng tên báo cáo nên dòng
        tên cột thật bị đẩy xuống thành dữ liệu, mọi dòng sau mang tên cột vô
        nghĩa 'Cột 3', 'Cột 8' — đoạn tra cứu được mà không đọc được."""
        from openpyxl import Workbook

        path = self.root / "bao-cao.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "BC Tháng 8"
        sheet.append(["BÁO CÁO CÔNG VIỆC THÁNG 08 NĂM 2025"])
        sheet.append([])
        sheet.append(["Số", "Ngày", "Phí dịch vụ"])
        sheet.append([1, "27/08/2025", 15600000])
        workbook.save(path)

        result = extract_text_with_metadata(path)

        self.assertIn("[Cột] Số | Ngày | Phí dịch vụ", result.text)
        self.assertIn("Phí dịch vụ: 15600000", result.text)
        self.assertNotIn("Cột 2:", result.text)
        # Dòng tên báo cáo vẫn được giữ làm ngữ cảnh, không biến thành tên cột.
        self.assertIn("BÁO CÁO CÔNG VIỆC THÁNG 08 NĂM 2025", result.text)
        self.assertNotIn("[Cột] BÁO CÁO", result.text)

    def test_xlsx_formula_without_cached_value_is_not_learned_as_formula(self):
        """data_only=True: đọc GIÁ TRỊ đã tính, không nhồi '=G14*8%' vào kho.
        File sinh bằng thư viện chưa có giá trị cache thì ô công thức bị bỏ,
        các ô giá trị thật vẫn được học."""
        from openpyxl import Workbook

        path = self.root / "cong-thuc.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Doanh thu", "Thuế"])
        sheet.append([26000000, "=A2*8%"])
        workbook.save(path)

        result = extract_text_with_metadata(path)

        self.assertIn("Doanh thu: 26000000", result.text)
        self.assertNotIn("=A2*8%", result.text)

    def test_xlsx_all_formula_file_falls_back_with_warning(self):
        """File toàn công thức (không có giá trị cache) thì đọc công thức thô
        kèm cảnh báo — còn hơn mất trắng nội dung."""
        from openpyxl import Workbook

        path = self.root / "toan-cong-thuc.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["=1+1", "=2+2"])
        workbook.save(path)

        result = extract_text_with_metadata(path)

        self.assertIn("=1+1", result.text)
        self.assertTrue(any("giá trị công thức" in warning
                            for warning in result.warnings))

    def test_context_header_marks_client_ownership(self):
        """Đoạn thuộc hồ sơ khách phải tự mang tên khách trong nội dung được
        embedding — 'Giấy đề nghị' của khách không được đọc y hệt hồ sơ HDS."""
        from app.ingest import ChunkPiece, apply_context_headers, context_header

        header = context_header("1. Giấy đề nghị", "filing", "CÔNG TY TNHH AGENT PRO")
        self.assertIn("hồ sơ khách hàng — CÔNG TY TNHH AGENT PRO", header)

        pieces = apply_context_headers(
            [ChunkPiece(content="Tổng số lao động (dự kiến): 02 người.")],
            "1. Giấy đề nghị", "filing", "CÔNG TY TNHH AGENT PRO")
        self.assertTrue(pieces[0].content.startswith("[Tài liệu: 1. Giấy đề nghị"))
        self.assertIn("Tổng số lao động", pieces[0].content)

        hr = apply_context_headers(
            [ChunkPiece(content="Hợp đồng lao động số 45HDLD-HDS.")],
            "45HDLD-HDS", "ho_so_ns", None)
        self.assertIn("hồ sơ nhân sự của công ty luật HDS", hr[0].content)

        # Đoạn LUẬT đã tự mang số hiệu văn bản — giữ nguyên, không gắn thêm.
        law = apply_context_headers(
            [ChunkPiece(content="[Bộ luật Lao động số 45/2019/QH14 — Điều 35]\nNội dung.")],
            "45_2019_QH14", "law", None)
        self.assertFalse(law[0].content.startswith("[Tài liệu:"))

    def test_corrupt_docx_has_stable_error_code(self):
        path = self.root / "hong.docx"
        path.write_bytes(b"not-a-docx")

        with self.assertRaises(ExtractionError) as caught:
            extract_text_with_metadata(path)

        self.assertEqual("invalid_docx", caught.exception.code)
        # API cũ vẫn không ném lỗi để giữ tương thích với caller hiện hữu.
        with redirect_stdout(io.StringIO()):
            self.assertEqual("", extract_text(path))

    def test_docx_reads_paragraphs_and_table_rows(self):
        from docx import Document

        path = self.root / "ho-so.docx"
        document = Document()
        document.add_heading("Nhân sự", level=1)
        document.add_paragraph("Thông tin hợp đồng lao động")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Nhân sự"
        table.cell(0, 1).text = "Trạng thái"
        table.cell(1, 0).text = "Lê Minh"
        table.cell(1, 1).text = "Còn hiệu lực"
        document.save(path)

        result = extract_text_with_metadata(path)

        self.assertEqual("python-docx", result.method)
        self.assertIn("Thông tin hợp đồng lao động", result.text)
        self.assertIn("Lê Minh | Còn hiệu lực", result.text)
        chunks = split_document_with_metadata(result, "other")
        self.assertEqual("Nhân sự", chunks[0].section_title)
        self.assertEqual("section:Nhân sự", chunks[0].source_locator)

    def test_pdf_page_markers_become_chunk_provenance(self):
        extraction = ExtractionResult(
            text="[Trang 1]\nNội dung trang thứ nhất.\n\n[Trang 2]\nNội dung trang thứ hai.",
            format="pdf", method="pdfplumber",
        )

        chunks = split_document_with_metadata(extraction, "other")

        self.assertEqual([1, 2], [chunk.page_number for chunk in chunks])
        self.assertEqual(["page:1", "page:2"], [chunk.source_locator for chunk in chunks])

    def test_empty_file_is_reported_explicitly(self):
        path = self.root / "rong.txt"
        path.write_bytes(b"")

        with self.assertRaises(ExtractionError) as caught:
            extract_text_with_metadata(path)

        self.assertEqual("empty_file", caught.exception.code)

    def test_drive_names_cannot_escape_destination(self):
        for unsafe in ("../secret.csv", "..", "a/b.xlsx", "a\\b.xlsx", "NUL.txt"):
            safe = safe_path_component(unsafe)
            self.assertNotIn("/", safe)
            self.assertNotIn("\\", safe)
            self.assertNotIn(safe, {"", ".", ".."})


class AutoLearnSafetyTests(unittest.TestCase):
    def test_review_is_default_and_legacy_env_remains_compatible(self):
        self.assertFalse(auto_approve_from_env({}))
        self.assertTrue(auto_approve_from_env({"AUTO_LEARN_AUTO_APPROVE": "1"}))
        self.assertFalse(auto_approve_from_env({"AUTO_LEARN_AUTO_APPROVE": "invalid"}))
        self.assertTrue(auto_approve_from_env({"AUTO_LEARN_REVIEW": "0"}))
        self.assertFalse(auto_approve_from_env({"AUTO_LEARN_REVIEW": "1"}))

    def test_matter_code_candidates_support_hds_numeric_convention(self):
        """Cây Drive thật đặt tên dự án bằng số ('1572. Thành lập…'), có khi kèm
        ngày phía trước ('160426. 1593. …'). Bot phải bắt được các mã đó, ưu
        tiên cụm số đứng sát tên, và không vớ số trong thân tên."""
        from app.auto_learn import _matter_code_candidates

        self.assertEqual(["M-2026-001"],
                         _matter_code_candidates("[M-2026-001] Tái cấu trúc vốn"))
        self.assertEqual(["1572"], _matter_code_candidates(
            "1572. Thành lập mới Công ty TNHH Học viện Ngôn ngữ Hoa Hạ"))
        self.assertEqual(["1593", "160426"], _matter_code_candidates(
            "160426. 1593. Thành lập mới CÔNG TY TNHH AGENT PRO"))
        self.assertEqual([], _matter_code_candidates(
            "Mua bán sáp nhập 2026 - CTCP Đại Hữu"))

    def test_google_native_file_has_stable_modified_time_fingerprint(self):
        info = {"modifiedTime": "2026-08-19T12:00:00.000Z"}
        self.assertEqual("gdrive-modified:2026-08-19T12:00:00.000Z",
                         drive_fingerprint(info))
        # File nhị phân giữ checksum cũ để không nạp lại một lần không cần thiết.
        self.assertEqual("abc", drive_fingerprint({"md5Checksum": "abc",
                                                   "modifiedTime": "later"}))

    def test_learn_one_writes_chunk_provenance_without_real_db_or_ollama(self):
        from openpyxl import Workbook

        class FakeCursor:
            def __init__(self):
                self.calls = []
                self.row = None

            def __enter__(self):
                return self

            def __exit__(self, *_args):
                return False

            def execute(self, sql, params=None):
                self.calls.append((sql, params))
                if "RETURNING id" in sql:
                    self.row = (42,)

            def fetchone(self):
                return self.row

        class FakeConnection:
            def __init__(self, cursor):
                self._cursor = cursor

            def __enter__(self):
                return self

            def __exit__(self, *_args):
                return False

            def cursor(self):
                return self._cursor

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "nhan-su.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Nhân sự"
            sheet.append(["Họ tên", "Trạng thái"])
            sheet.append(["Phạm An", "Còn hợp đồng"])
            workbook.save(path)

            cursor = FakeCursor()
            diagnostics = {}
            labels = {"doc_type": "ho_so_ns", "access_level": "internal",
                      "client_id": None, "department_id": None, "matter_id": None}
            with patch.object(auto_learn, "embed", return_value=[[0.1, 0.2]]), \
                    patch.object(auto_learn, "summarize", return_value="Tóm tắt"), \
                    patch.object(auto_learn.db, "session", return_value=FakeConnection(cursor)), \
                    patch.object(auto_learn.db, "audit"):
                with redirect_stdout(io.StringIO()):
                    learned = auto_learn.learn_one(
                        path, labels, "drive-id", "checksum", diagnostics=diagnostics)

        self.assertTrue(learned)
        chunk_call = next(call for call in cursor.calls if "INSERT INTO chunks" in call[0])
        self.assertIn("page_number", chunk_call[0])
        self.assertEqual("Nhân sự", chunk_call[1][4])
        self.assertEqual("sheet:Nhân sự;rows:2-2", chunk_call[1][5])
        self.assertEqual("xlsx", diagnostics["method"])


class SmartChunkingTests(unittest.TestCase):
    """Cắt đoạn theo ngữ cảnh: không cắt giữa câu, ngắt ở chỗ chuyển ý."""

    def test_short_document_stays_whole(self):
        """Sơ yếu lý lịch một trang không có lý do gì để bị xé nhỏ."""
        text = "Họ tên: Bạc Thị Mai. Chức danh Trưởng phòng. Vào làm năm 2022."
        self.assertEqual(len(chunk_generic(text)), 1)

    def test_never_cuts_mid_sentence(self):
        long_text = " ".join(
            f"Câu số {i} nói về một nội dung dài vừa phải để kiểm tra việc cắt đoạn."
            for i in range(200))
        for piece in chunk_generic(long_text, target_words=60):
            # Mọi đoạn phải bắt đầu bằng đầu một câu và kết thúc bằng dấu câu.
            self.assertTrue(piece.strip().startswith("Câu số"), piece[:60])
            self.assertTrue(piece.strip().endswith("."), piece[-60:])

    def test_abbreviation_does_not_end_sentence(self):
        """'Nghị định số 01/2021/NĐ-CP.' không được tách làm hai đơn vị."""
        units = _split_units("Áp dụng theo NĐ. 01/2021 của Chính phủ. Hết.")
        self.assertTrue(any("NĐ. 01/2021" in u for u in units),
                        f"Bị tách sai: {units}")

    def test_heading_starts_new_chunk(self):
        text = ("THÔNG TIN CHUNG\n"
                + "Nội dung phần một nói về hợp đồng và các bên tham gia. " * 12
                + "\nĐIỀU KHOẢN THANH TOÁN\n"
                + "Nội dung phần hai nói về tiền và thời hạn trả. " * 12)
        pieces = chunk_generic(text, target_words=80)
        # Hai tiêu đề phải nằm ở ĐẦU hai đoạn khác nhau, không lẫn vào giữa.
        starts = [p.split("\n")[0][:30] for p in pieces]
        self.assertTrue(any(s.startswith("THÔNG TIN CHUNG") for s in starts), starts)
        self.assertTrue(any(s.startswith("ĐIỀU KHOẢN THANH TOÁN") for s in starts), starts)

    def test_table_rows_are_not_split(self):
        text = "\n".join(f"[Dòng {i}] Tên: Người {i} | Chức danh: Nhân viên"
                         for i in range(1, 40))
        for piece in chunk_generic(text, target_words=40):
            # Không đoạn nào được kết thúc giữa chừng một dòng bảng.
            self.assertFalse(piece.rstrip().endswith("|"), piece[-40:])

    def test_chunk_sizes_follow_content_not_a_fixed_ruler(self):
        """Mục ngắn giữ trọn một đoạn; mục dài mới bị tách — nên độ dài KHÔNG đều.

        Đây là điểm khác biệt so với bản cắt cứng 320 từ: kích thước đoạn do nội
        dung quyết định, không do một con số định sẵn.
        """
        text = ("MỤC A\n" + "Nội dung ngắn. " * 8
                + "\nMỤC B\n" + "Nội dung dài hơn nhiều lần so với mục trước. " * 30)
        sizes = [len(p.split()) for p in chunk_generic(text, target_words=40)]
        self.assertGreater(len(sizes), 2)
        # Mục A (~26 từ) phải nhỏ hơn hẳn ngân sách, không bị độn cho đủ 40.
        self.assertLess(min(sizes), 35, f"Mục ngắn bị độn: {sizes}")
        self.assertGreater(max(sizes) - min(sizes), 5, f"Kích thước quá đều: {sizes}")


class LawChunkingTests(unittest.TestCase):
    """Văn bản quy phạm phải cắt theo Điều và mang theo số hiệu văn bản.

    Không có số hiệu thì mọi đoạn đều là 'Điều 5' trơ trọi — bot dẫn nguồn kiểu
    'theo quy định pháp luật', không dùng được trong hành nghề luật.
    """

    LAW = """BỘ LUẬT LAO ĐỘNG
Số: 45/2019/QH14

Chương I
NHỮNG QUY ĐỊNH CHUNG

Điều 1. Phạm vi điều chỉnh
Bộ luật này quy định tiêu chuẩn lao động và quan hệ lao động.

Điều 2. Đối tượng áp dụng
Người lao động và người sử dụng lao động trên lãnh thổ Việt Nam.

Chương II
VIỆC LÀM

Điều 35. Quyền đơn phương chấm dứt hợp đồng
Người lao động có quyền đơn phương chấm dứt hợp đồng lao động.
"""

    def test_document_number_is_read_from_header(self):
        self.assertIn("45/2019/QH14", document_citation(self.LAW))

    def test_reference_in_body_does_not_override_header(self):
        """Số hiệu phải lấy ở phần đầu; thân bài đầy số hiệu do dẫn chiếu."""
        text = self.LAW + "\nĐiều 99. Dẫn chiếu\nTheo Nghị định 145/2020/NĐ-CP.\n"
        self.assertIn("45/2019/QH14", document_citation(text))

    def test_each_article_becomes_one_chunk(self):
        pieces = chunk_law_structured(self.LAW)
        self.assertEqual(len(pieces), 3)

    def test_chunk_carries_full_citation_path(self):
        pieces = chunk_law_structured(self.LAW)
        last = pieces[-1]
        # Đoạn cuối phải biết mình là Điều 35, Chương II, của văn bản nào.
        self.assertIn("Điều 35", last.section_title)
        self.assertIn("Chương II", last.section_title)
        self.assertIn("45/2019/QH14", last.section_title)
        self.assertEqual("dieu:35", last.source_locator)

    def test_citation_is_inside_content_not_only_metadata(self):
        """Đoạn vào prompt dưới dạng văn bản thuần — model chỉ đọc được nội dung."""
        pieces = chunk_law_structured(self.LAW)
        self.assertIn("45/2019/QH14", pieces[-1].content)

    def test_text_without_articles_falls_back_to_generic(self):
        pieces = chunk_law_structured("Một văn bản không có điều khoản nào cả.")
        self.assertTrue(pieces)
        self.assertEqual("document", pieces[0].source_locator)


if __name__ == "__main__":
    unittest.main()
